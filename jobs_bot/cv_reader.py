from __future__ import annotations

from pathlib import Path

from docx import Document


def read_docx_text(path: str | Path) -> str:
    """Extract plain text from a .docx CV.

    Extraction is intentionally simple and deterministic:
    - paragraph text (in order)
    - table cell text (in order)

    Raises:
        ValueError: if the file cannot be parsed as a DOCX document.
    """
    p = Path(path)
    try:
        doc = Document(str(p))
    except Exception as exc:  # pragma: no cover
        raise ValueError(f"Cannot parse DOCX file: {p}") from exc

    parts: list[str] = []

    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    parts.append(txt)

    return "\n".join(parts).strip()
